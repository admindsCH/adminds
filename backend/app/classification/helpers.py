"""Helpers to extract text from uploaded files using LiteParse and GPT vision."""

from __future__ import annotations

import base64
import io
import tempfile
from pathlib import PurePath

from docx import Document
from langchain_core.messages import HumanMessage, SystemMessage
from liteparse import LiteParse
from loguru import logger

_SUPPORTED_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".doc",
    ".jpg",
    ".jpeg",
    ".png",
    ".tiff",
    ".bmp",
    ".heic",
}

_IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".tiff", ".bmp", ".heic"}

_parser = LiteParse()


def is_image_file(filename: str) -> bool:
    """Return True if the filename has an image extension."""
    return PurePath(filename).suffix.lower() in _IMAGE_EXTENSIONS


def _convert_heic_to_jpeg(file_bytes: bytes) -> bytes:
    """Convert HEIC image bytes to JPEG bytes."""
    from pillow_heif import register_heif_opener
    from PIL import Image

    register_heif_opener()
    img = Image.open(io.BytesIO(file_bytes))
    buf = io.BytesIO()
    img.convert("RGB").save(buf, format="JPEG", quality=90)
    return buf.getvalue()


def _prepare_image_for_vision(filename: str, file_bytes: bytes) -> tuple[str, str]:
    """Return (mime_type, base64_data) ready for GPT vision.

    Converts HEIC to JPEG; other formats are passed through.
    """
    ext = PurePath(filename).suffix.lower()
    if ext == ".heic":
        file_bytes = _convert_heic_to_jpeg(file_bytes)
        mime_type = "image/jpeg"
    else:
        mime_map = {
            ".jpg": "image/jpeg",
            ".jpeg": "image/jpeg",
            ".png": "image/png",
            ".tiff": "image/tiff",
            ".bmp": "image/bmp",
        }
        mime_type = mime_map.get(ext, "image/png")
    return mime_type, base64.b64encode(file_bytes).decode("utf-8")


_VISION_TRANSCRIPTION_PROMPT = """\
Tu es un expert en lecture de documents médicaux manuscrits en français.

Tu reçois l'image d'un document manuscrit (notes cliniques, ordonnance, etc.).
Transcris INTÉGRALEMENT le texte manuscrit visible dans l'image, en français.

Règles:
- Transcris fidèlement ce qui est écrit, même si l'écriture est difficile à lire
- Préserve la structure (paragraphes, listes, tirets) autant que possible
- Si un mot est illisible, écris [illisible]
- N'invente RIEN — ne transcris que ce que tu vois
- Inclus les dates, noms, dosages de médicaments tels quels"""


async def extract_text_vision(filename: str, file_bytes: bytes) -> str:
    """Extract text from an image using GPT vision (for handwritten documents).

    Encodes the image as base64, sends it to GPT vision with a transcription
    prompt, and returns the transcribed text.

    Args:
        filename: Original filename (used for logging).
        file_bytes: Raw image content.

    Returns:
        Transcribed text string.
    """
    from app.services.azure_openai import ainvoke_throttled, get_model

    mime_type, b64 = _prepare_image_for_vision(filename, file_bytes)

    logger.info(f"Extracting text via GPT vision from '{filename}'...")
    model = get_model()
    result = await ainvoke_throttled(
        model,
        [
            SystemMessage(content=_VISION_TRANSCRIPTION_PROMPT),
            HumanMessage(
                content=[
                    {
                        "type": "image_url",
                        "image_url": {
                            "url": f"data:{mime_type};base64,{b64}",
                        },
                    },
                ]
            ),
        ],
    )
    text = result.content
    logger.info(f"Vision extracted {len(text)} chars from '{filename}'")
    return text


def _extract_text_docx(file_bytes: bytes) -> str:
    """Extract text from a .docx file using python-docx.

    Reads all paragraphs and table cells, preserving document order.
    """
    doc = Document(io.BytesIO(file_bytes))
    parts: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_texts:
                parts.append(" | ".join(row_texts))

    return "\n\n".join(parts)


def extract_text(filename: str, file_bytes: bytes) -> str:
    """Extract text from any supported file.

    Uses python-docx for .docx files, LiteParse for everything else
    (PDF, images with OCR, etc.).

    Args:
        filename: Original filename (used for extension detection).
        file_bytes: Raw file content.

    Returns:
        Extracted text string.
    """
    ext = PurePath(filename).suffix.lower()

    if ext not in _SUPPORTED_EXTENSIONS:
        logger.warning(f"Unsupported file type: {filename}")
        return f"[Fichier non supporté: {filename}]"

    # .docx: use python-docx directly (faster, more reliable)
    if ext == ".docx":
        logger.info(f"Extracting text from '{filename}' via python-docx...")
        try:
            text = _extract_text_docx(file_bytes)
            logger.info(f"Extracted {len(text)} chars from '{filename}'")
            return text
        except Exception:
            logger.exception(f"python-docx failed for '{filename}'")
            return f"[Erreur d'extraction: {filename}]"

    # Everything else: LiteParse (PDF, .doc, images with OCR)
    with tempfile.NamedTemporaryFile(suffix=ext, delete=True) as tmp:
        tmp.write(file_bytes)
        tmp.flush()

        logger.info(f"Extracting text from '{filename}' via LiteParse...")
        try:
            result = _parser.parse(tmp.name, ocr_language="fra")
        except Exception:
            logger.exception(f"LiteParse failed for '{filename}'")
            return f"[Erreur d'extraction: {filename}]"

    logger.info(
        f"Extracted {len(result.text)} chars from '{filename}' "
        f"({len(result.pages)} page(s))"
    )
    return result.text
