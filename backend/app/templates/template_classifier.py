"""Template classifier — auto-detect metadata from an uploaded template.

Reads the document content and uses GPT structured output to determine:
- name: clean display name (e.g. "Rapport AI Fribourg")
- description: one-line description
- category: rapport-ai | rapport-medical | rapport-assurance | rapport-perte-gain
- canton: fribourg | geneve | all
- insurance: matched against known insurances, or empty
- page_count: estimated from content
"""

from __future__ import annotations

import io

import fitz  # PyMuPDF
from docx import Document
from langchain_core.messages import HumanMessage, SystemMessage
from loguru import logger

from app.services.azure_openai import ainvoke_throttled, get_model
from app.templates.schemas import TemplateClassification

# Known insurance names → IDs
_KNOWN_INSURANCES = {
    "suva": "ins-suva",
    "css": "ins-css",
    "helsana": "ins-helsana",
    "visana": "ins-visana",
    "groupe mutuel": "ins-gm",
    "ai": "ins-ai",
    "ai fédérale": "ins-ai",
    "assurance invalidité": "ins-ai",
    "asa": "ins-asa",
    "svv": "ins-asa",
}

_CLASSIFIER_PROMPT = """\
Tu es un expert en formulaires médicaux suisses pour les assurances sociales \
(AI, LAA, LAMal, LPP, assurance militaire).

On te donne le contenu textuel d'un formulaire/rapport médical vide (template). \
Analyse-le et extrais ses métadonnées."""


async def classify_template(
    file_bytes: bytes, *, is_pdf: bool = False
) -> dict[str, str]:
    """Classify a template document and return metadata.

    Args:
        file_bytes: .docx or .pdf file bytes.
        is_pdf: True if the file is a PDF.

    Returns:
        Dict with keys: name, description, category, canton, insurance_id,
        page_count, estimated_minutes.
    """
    if is_pdf:
        content = _extract_pdf_text(file_bytes, max_chars=3000)
    else:
        content = _extract_docx_text(file_bytes, max_chars=3000)

    if not content.strip():
        logger.warning("Template has no text content, using defaults")
        return _default_metadata()

    logger.info(f"Classifying template ({len(content)} chars of content)")

    structured = get_model().with_structured_output(TemplateClassification)
    result: TemplateClassification = await ainvoke_throttled(
        structured,
        [
            SystemMessage(content=_CLASSIFIER_PROMPT),
            HumanMessage(content=f"Contenu du formulaire:\n\n{content}"),
        ],
    )

    # Map insurance name to ID
    insurance_lower = result.insurance.lower().strip()
    insurance_id = ""
    for known_name, known_id in _KNOWN_INSURANCES.items():
        if known_name in insurance_lower or insurance_lower in known_name:
            insurance_id = known_id
            break

    metadata = {
        "name": result.name,
        "description": result.description,
        "category": result.category,
        "canton": result.canton,
        "insurance_id": insurance_id,
        "insurance_name": result.insurance,
        "page_count": str(result.page_count),
        "estimated_minutes": str(max(2, result.page_count)),
        "is_official": "false",
    }

    logger.info(
        f"Template classified: name='{metadata['name']}', "
        f"category={metadata['category']}, canton={metadata['canton']}"
    )

    return metadata


def _extract_pdf_text(pdf_bytes: bytes, max_chars: int = 3000) -> str:
    """Extract text from a PDF using PyMuPDF."""
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    parts: list[str] = []
    total = 0

    for page in doc:
        text = page.get_text().strip()
        if text:
            parts.append(text)
            total += len(text)
            if total >= max_chars:
                break

    doc.close()
    return "\n".join(parts)[:max_chars]


def _extract_docx_text(docx_bytes: bytes, max_chars: int = 3000) -> str:
    """Extract text from a .docx file."""
    from app.templates.docx_compat import normalize_docx_bytes

    doc = Document(io.BytesIO(normalize_docx_bytes(docx_bytes)))
    parts: list[str] = []
    total = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
            total += len(text)
            if total >= max_chars:
                break

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                line = " | ".join(cells)
                parts.append(line)
                total += len(line)
                if total >= max_chars:
                    break
        if total >= max_chars:
            break

    return "\n".join(parts)[:max_chars]


def _default_metadata() -> dict[str, str]:
    return {
        "name": "Formulaire importé",
        "description": "Formulaire médical importé",
        "category": "rapport-medical",
        "canton": "all",
        "insurance_id": "",
        "insurance_name": "",
        "page_count": "1",
        "estimated_minutes": "3",
        "is_official": "false",
    }
